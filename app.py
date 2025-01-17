import os
# from OPENAI import client
import shutil
import tempfile
import hmac
import openai
from langchain.document_loaders import SQLDatabaseLoader
from langchain.vectorstores import Chroma
from langchain.llms import OpenAI
from langchain.chains import RetrievalQA
from fastapi import FastAPI
from dotenv import load_dotenv
from langchain.docstore.document import Document
import snowflake.connector
from snowflake import connector
import pandas as pd
from docx import Document
import streamlit as st
from fastapi.responses import FileResponse
from io import BytesIO
import re

# Load environment variables from .env file
from dotenv import load_dotenv
load_dotenv()

api_key = os.environ.get("OPENAI_API_KEY")
client = openai.Client(api_key=api_key)

# Set up Snowflake connection
username = os.getenv("SNOWFLAKE_USERNAME")
password = os.getenv("SNOWFLAKE_PASSWORD")
account_name = os.getenv("SNOWFLAKE_ACCOUNT_NAME")
warehouse = os.getenv("SNOWFLAKE_WAREHOUSE")
database = os.getenv("SNOWFLAKE_DATABASE")
schema = os.getenv("SNOWFLAKE_SCHEMA")

ctx = snowflake.connector.connect(user=username, password=password, account=account_name,
                                   warehouse=warehouse, database=database, schema=schema, protocol='https')


connection_string = "snowflake://{username}:{password}@{account_name}/?"
connection_string += "warehouse={warehouse}&db={database}&schema={schema}"


conn = connector.connect(
    user=username,
    password=password,
    account=account_name,
    warehouse=warehouse,
    database=database,
    schema=schema
)
snowflake_loader = SQLDatabaseLoader(
    conn.cursor(),
    db=database
)

def execute_query(cursor, query):
    cursor.execute(query)
    return cursor.fetchall()


# Define Streamlit app
def main_app():
    st.title("Medical Summary Generator")

    # Get patient ID from user input
    patient_id = st.text_input("Enter Patient ID:")

    if st.button("Generate Medical Summary"):
        with st.spinner(
            'Analyzing Patient Reports...'
        ):
            report_buffer = generate_medical_summary(patient_id)

            st.download_button(
                label="Download Report",
                data=report_buffer.getvalue(),
                file_name=f"{patient_id} Health Insurance.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        # Add your application's main functionality here or redirect to another app page



# Function to load data and create the medical summary
def generate_medical_summary(patient_id):
    st.spinner(
        'Analyzing Patient Reports..'
    )
    docs_df_report_image, docs_df_report_lab, docs_df_report_notes = load_data_and_create_vector_db(patient_id)

    #                                             """
    initial_prompt = """
    You are a medical professional, You need to compose a detailed summary report based on the provided doctor's notes, image reports, and lab reports for a patient. The report is intended for the patient's primary care physician, and should be written in clear, concise language that they can easily understand.
                            """

    prompt = """
    #FYI
    who you are?
    You are a medical professional, You need to compose a detailed summary report based on the provided doctor's notes, image reports, and lab reports for a patient. The report is intended for the patient's primary care physician, and should be written in clear, concise language that they can easily understand.
    you have already provided with the data of doctor notes, image notes and lab reports in our previous conversation itself please use that as the data here
    what you have to do?
    You have to go through the following instruction and notes below to form a neat summary with the output in the bottom

    #Instruction
    The summary should highlight each identified medical Category, categorized by appropriate medical terms and organized under relevant sections such as:

    * Cardiovascular system
    * Cancer screening
    * Body composition
    * Additional laboratory testing

    For each description under each category, please provide an in-depth analysis that includes:

    * Specific values and units of measurement
    * Comparisons to normal ranges
    * Potential implications for the patient's health
    * The patient's concerns or questions, if applicable then only mention the point

    It's crucial to cite the source of data (doctor's note, image report, lab report) for each result included in the summary report.

    The report should cover at least 10 medical descriptions and follow this format:

    Additional notes regarding the medical terms listed in the provided data:

    Summary of Medical Testing
    ----------------------------

    (Brief overview of the findings from the medical report.)

    Over the course of your Personalized Health Assessment, you had comprehensive testing for cardiovascular disease, cancer, and other medical conditions. Thankfully, your results were very reassuring. Below is a list of your tests along with a brief summary of the results. The full results are included in the Appendix of this document.

    Cardiovascular Testing – Normal, aside from mildly elevated lipids and suboptimal endurance.
    ●	Boston Heart Labs: As we discussed, your cholesterol parameters are not ideal. For now, the treatment should be a healthy lifestyle, in particular:
        ○	Increasing aerobic activity to three times per week.
        ○	Replace dietary sugars with complex carbohydrates from unprocessed vegetables.
        ○	Replace dietary saturated fats with unsaturated fats (eg, butter→olive oil, meat→wild fish)
    We will repeat a standard lipid panel at your 3-month follow-up visit. (See appendix, pages 10-16)
    ●	Electrocardiogram (ECG): Normal. The computer noted “Early Repolarization,” which is not concerning. (Page 17)
    ●	Ultrasound Carotid Doppler: Normal – the arteries in your neck do not show any signs of plaque. (Pages 18-19)
    ●	CT Coronary Calcium Score = zero – excellent! (Pages 20-21)
    ●	Resting Echocardiogram: Normal – your heart is functioning normally at rest. The “trace regurgitation” of your mitral, tricuspid, and pulmonic valves is considered normal. (Pages 22-24).
    ●	Stress Echocardiogram: Normal – your heart responds normally to exercise. However, your exercise capacity was not as vigorous as we would like. Again, the plan is to increase the frequency of your aerobic exercise (Pages 25-28)
    
    Cancer Screening – All Negative, except for elevated liver fat, and an incidental finding noted below, which we will reassess:
    ●	Grail Galleri Multi-Cancer Blood Test: Negative – no signal detected for the more than 20 cancers tested. (Pages 29-35)
    ●	Thyroid Sonogram: Normal – No nodules. (Page 36)
    ●	CT Scan of the Chest: There was an incidental finding of “ground glass opacities,” and lower lung “nodularity,” which was felt to be either inflammatory or infectious. Since Autumn had a similar finding, I suspect you both had the same viral respiratory infection, which left similar imprints on your Chest CTs. We will repeat this scan at your three-month follow up to reevaluate. (Pages 37)
    ●	MRI of the Abdomen and Pelvis: Thankfully, no concerning masses were seen. You did have mildly elevated liver fat. This was likely the cause of your previously elevated liver function tests, which have since normalized. Fatty liver disease generally tracks with weight and typically improves with weight loss. The two most specific dietary contributors to fatty liver are sugar and alcohol. Since you rarely drink alcohol, I think the simplest measure you could implement to reduce your liver fat would be to reduce your dietary sugar. I recommend you work to address this, and we can reassess liver imaging in six months. (Pages 38-39)
    ●	MRI of the Prostate: Normal – no concerning lesions. (Pages 40-41)
    
    Body Composition:
    ●	DEXA Bone Scan: Good - your bones look strong overall. A couple of vertebrae were marked as being on the weaker side. I would not put too much emphasis on this, but just another reason to continue your new fitness regimen and strengthen your core. (Pages 42-44,47-49)
    ●	DEXA Body Composition: Normal. Your body fat percentage is 22.2. (Pages 45-46)
    ●	AMRA Body Composition: Excellent muscle volume based upon your thighs. Very low fat around your organs (visceral), beneath your skin (subcutaneous), and in your muscles. Again, your liver fat came out high. (Pages 50-51)
    
    Additional Laboratory Testing: (Page 52-65)
    ●	Your initial labs from February were all normal, except for a few minor findings:
        ○	Elevated liver enzymes (AST, ALT), but these normalized upon repeating.
        ○	Mildly elevated LDL cholesterol; the plan is discussed above.
        ○	Mildly elevated thyroid hormone Free T4, but this was fine upon repeating.
        ○	Your Free Testosterone was flagged as mildly elevated, but the repeat was fine. Note, that these were different labs, with different reference values.
        ○	You are immune to Hepatitis A and B and you are negative for Hepatitis C.
    ●	Repeat labs from March were excellent and showed normalization of the above abnormal findings. 


    Please ensure to:

    * Strictly follow the above mentioned output format
    * Identify and categorize each medical Description
    * Document the patient's concerns or questions, Don't mention it if not applicable
    * Cite the source of data for each result
    * Include all necessary medical tests in the report
    * Be absolute specific when you advise suggest or recommend something, ex -  dietary adjustment - Replace dietary saturated fats with unsaturated fats (eg, butter→olive oil, meat→wild fish), Instead of saying follow-up test is required write follow-up test in 3 to 6 months is required.

    The desired length of the report is 2-3 pages, so be focused in your analysis.

    #follow the output format
    Note - 1. Each pointers must be at least 200 words. You also have to give an explanation for your response. 
           2. If you are suggesting some adjustments always give some examples for it. for ex - dietary adjustment - Replace dietary saturated fats with unsaturated fats (eg, butter→olive oil, meat→wild fish)
           3. If you are recommending or advising something to the patient be absolute specific. for ex - Instead of saying follow-up test is required write follow-up test in 3 to 6 months is required.
           4. Apart from description in each subtopic you need to give three suggestions or tips to keep the patient health very strong, if condition is worse give the patient what will be the cause for the worst condition.
           5. Do not give any warning notes in the output and also don't make any initial sentences or content like "Given the constraints of this task and the nature of the provided data, a comprehensive and detailed medical report as described cannot be generated.". Strictly stick to the output format.
           

    output format:
    Summary of Medical Testing
    (Write 100 words about over all summary or the health condition of the patient)

    Overview of the findings
    Cardiovascular System
        ● Boston Heart Labs            
            ○ Description(200-300 words) (specific values, units, comparison to normal range) AND Potential implications for patient's health AND Source (doctor's note, image report, lab report) - In a single paragraph. \n
            ○ Suggestion - If the description about the patient condition is worse prioritize on suggesting the cause and recovery tips, else only suggest about diet, excersises, tips for keeping your health good on the current topic. \n
            Evaluation - You must pick only option from these options - Normal / Elevated / Low / Severe / Acute based on which relates best with the description of the current topic. \n
        ● Electrocardiogram (ECG)
            ○ Description(200-300 words) (specific values, units, comparison to normal range) AND Potential implications for patient's health AND Source (doctor's note, image report, lab report)- In a single paragraph. \n
            ○ Suggestion - If the description about the patient condition is worse prioritize on suggesting the cause and recovery tips, else only suggest about diet, excersises, tips for keeping your health good on the current topic. \n
            Evaluation - You must pick only option from these options - Normal / Elevated / Low / Severe / Acute based on which relates best with the description of the current topic. \n
        ● Ultrasound Carotid Doppler
            ○ Description(200-300 words) (specific values, units, comparison to normal range) AND Potential implications for patient's health AND Source (doctor's note, image report, lab report)- In a single paragraph. \n
            ○ Suggestion - If the description about the patient condition is worse prioritize on suggesting the cause and recovery tips, else only suggest about diet, excersises, tips for keeping your health good on the current topic. \n
            Evaluation - You must pick only option from these options - Normal / Elevated / Low / Severe / Acute based on which relates best with the description of the current topic. \n
        ● CT Coronary Calcium Score
            ○ Description(200-300 words) (specific values, units, comparison to normal range) AND Potential implications for patient's health AND Source (doctor's note, image report, lab report)- In a single paragraph. \n
            ○ Suggestion - If the description about the patient condition is worse prioritize on suggesting the cause and recovery tips, else only suggest about diet, excersises, tips for keeping your health good on the current topic. \n
            Evaluation - You must pick only option from these options - Normal / Elevated / Low / Severe / Acute based on which relates best with the description of the current topic.\n
        ● Resting Echocardiogram
            ○ Description(200-300 words) (specific values, units, comparison to normal range) AND Potential implications for patient's health AND Source (doctor's note, image report, lab report)- In a single paragraph.\n
            ○ Suggestion - If the description about the patient condition is worse prioritize on suggesting the cause and recovery tips, else only suggest about diet, excersises, tips for keeping your health good on the current topic. \n
            Evaluation - You must pick only option from these options - Normal / Elevated / Low / Severe / Acute based on which relates best with the description of the current topic.\n
        ● Stress Echocardiogram
            ○ Description (specific values, units, comparison to normal range) AND Potential implications for patient's health AND Source (doctor's note, image report, lab report)- In a single paragraph.\n
            ○ Suggestion - If the description about the patient condition is worse prioritize on suggesting the cause and recovery tips, else only suggest about diet, excersises, tips for keeping your health good on the current topic. \n
            Evaluation - You must pick only option from these options - Normal / Elevated / Low / Severe / Acute based on which relates best with the description of the current topic.\n
    ...
    Cancer Screening
        ● Grail Galleri Multi
            ○ Description (specific values, units, comparison to normal range) AND Potential implications for patient's health AND Source (doctor's note, image report, lab report)- In a single paragraph.\n
            ○ Suggestion - If the description about the patient condition is worse prioritize on suggesting the cause and recovery tips, else only suggest about diet, excersises, tips for keeping your health good on the current topic. \n
            Evaluation - You must pick only option from these options - Normal / Elevated / Low / Severe / Acute based on which relates best with the description of the current topic.\n
        ● Thyroid Sonogram
            ○ Description (specific values, units, comparison to normal range) AND Potential implications for patient's health AND Source (doctor's note, image report, lab report)- In a single paragraph.\n
            ○ Suggestion - If the description about the patient condition is worse prioritize on suggesting the cause and recovery tips, else only suggest about diet, excersises, tips for keeping your health good on the current topic. \n
            Evaluation - You must pick only option from these options - Normal / Elevated / Low / Severe / Acute based on which relates best with the description of the current topic.\n
        ● CT Scan of the Chest
            ○ Description (specific values, units, comparison to normal range) AND Potential implications for patient's health AND Source (doctor's note, image report, lab report)- In a single paragraph.\n
            ○ Suggestion - If the description about the patient condition is worse prioritize on suggesting the cause and recovery tips, else only suggest about diet, excersises, tips for keeping your health good on the current topic. \n
            Evaluation - You must pick only option from these options - Normal / Elevated / Low / Severe / Acute based on which relates best with the description of the current topic.\n
        ● MRI of the Abdomen and Pelvis
            ○ Description (specific values, units, comparison to normal range) AND Potential implications for patient's health AND Source (doctor's note, image report, lab report)- In a single paragraph.\n
            ○ Suggestion - If the description about the patient condition is worse prioritize on suggesting the cause and recovery tips, else only suggest about diet, excersises, tips for keeping your health good on the current topic. \n
            Evaluation - You must pick only option from these options - Normal / Elevated / Low / Severe / Acute based on which relates best with the description of the current topic.\n
        ● MRI of the Prostate
            ○ Description (specific values, units, comparison to normal range) AND Potential implications for patient's health AND Source (doctor's note, image report, lab report)- In a single paragraph.\n
            ○ Suggestion - If the description about the patient condition is worse prioritize on suggesting the cause and recovery tips, else only suggest about diet, excersises, tips for keeping your health good on the current topic. \n
            Evaluation - You must pick only option from these options - Normal / Elevated / Low / Severe / Acute based on which relates best with the description of the current topic.\n

    ...
    Body Composition
        ● DEXA Bone Scan
            ○ Description (specific values, units, comparison to normal range) AND Potential implications for patient's health AND Source (doctor's note, image report, lab report)- In a single paragraph.\n
            ○ Suggestion - If the description about the patient condition is worse prioritize on suggesting the cause and recovery tips, else only suggest about diet, excersises, tips for keeping your health good on the current topic. \n
            Evaluation - You must pick only option from these options - Normal / Elevated / Low / Severe / Acute based on which relates best with the description of the current topic.\n
        ● DEXA Body Composition
            ○ Description (specific values, units, comparison to normal range) AND Potential implications for patient's health AND Source (doctor's note, image report, lab report)- In a single paragraph.\n
            ○ Suggestion - If the description about the patient condition is worse prioritize on suggesting the cause and recovery tips, else only suggest about diet, excersises, tips for keeping your health good on the current topic. \n
            Evaluation - You must pick only option from these options - Normal / Elevated / Low / Severe / Acute based on which relates best with the description of the current topic.\n
        ● AMRA Body Composition
            ○ Description (specific values, units, comparison to normal range) AND Potential implications for patient's health AND Source (doctor's note, image report, lab report)- In a single paragraph.\n
            ○ Suggestion - If the description about the patient condition is worse prioritize on suggesting the cause and recovery tips, else only suggest about diet, excersises, tips for keeping your health good on the current topic. \n
            Evaluation - You must pick only option from these options - Normal / Elevated / Low / Severe / Acute based on which relates best with the description of the current topic.\n

    ...
    Additional Laboratory Testing
        ● Category 1
            ○ Description (specific values, units, comparison to normal range) AND Potential implications for patient's health AND Source (doctor's note, image report, lab report)\n
            ○ Suggestion - If the description about the patient condition is worse prioritize on suggesting the cause and recovery tips, else only suggest about diet, excersises, tips for keeping your health good on the current topic. \n
            Evaluation - You must pick only option from these options - Normal / Elevated / Low / Severe / Acute based on which relates best with the description of the current topic.\n
        ● Category 2
            ○ Description (specific values, units, comparison to normal range) AND Potential implications for patient's health AND Source (doctor's note, image report, lab report)\n
            ○ Suggestion - If the description about the patient condition is worse prioritize on suggesting the cause and recovery tips, else only suggest about diet, excersises, tips for keeping your health good on the current topic. \n
            Evaluation - You must pick only option from these options - Normal / Elevated / Low / Severe / Acute based on which relates best with the description of the current topic.\n
    ...
    
    (Strictly follow the above mentioned output format)
    (At least 10 Descriptions total across the sections, following the substructure provided for description, implications, source, and patient concerns)


    Note:
    Remember that you are a medical professional and you have a data of patient's full body health check up result and you only job is to summarize the data in the given ouptut format.
    Make sure to start with summary of medical testing as mentioned in the output format, nothing else.
    """

    messages = [
        {"role": "user", "content": initial_prompt},
        {"role": "system",
         "content": "I will be acting as a medical professional for you and will help you out giving the summary with the provided data."},
        {"role": "user", "content": f"Doctors Notes: {docs_df_report_notes}"},
        {"role": "user", "content": f"Image Reports: {docs_df_report_image}"},
        {"role": "user", "content": f"Lab Reports: {docs_df_report_lab}"},
        {"role": "user",
         "content": prompt}
        # {"role": "user", "content": f"The ID of the patient is: {patient_id}."}
    ]

    response = client.chat.completions.create(
        model="gpt-4-0125-preview",
        temperature=0,
        messages=messages,
        max_tokens=4000,
        top_p=0.0001
    )
    response_text = response.choices[0].message.content
    print(response.choices[0].message.content)

    st.success(response.choices[0].message.content)

    doc = Document()

    lines = response_text.split("\n")

    # Iterate through the lines and add them to the document
    for line in lines:
        # Check if the line starts with '<document_content>'
        if line.startswith("<document_content>"):
            doc.add_heading(line.strip("<document_content>").strip(), level=0)
            continue

        # Check if the line starts with '###'
        if line.startswith("###"):
            doc.add_heading(line.strip("#").strip(), level=2)
            continue

        # Check if the line starts with '**'
        if line.startswith("**"):
            doc.add_heading(line.strip("**").strip(), level=1)
            continue

        # Check if the line starts with '-'
        if line.startswith("-"):
            # Split the line into parts
            parts = line.split("**")

            # Add a new paragraph for the bullet point
            p = doc.add_paragraph(style="List Bullet")

            # Add the first part as regular text
            p.add_run(parts[0].strip("- "))

            # Add the remaining parts as bold text
            for part in parts[1:]:
                bold_part = part.strip("*")
                p.add_run(bold_part).bold = True
            continue
        line = re.sub(r'\\\*', ' ', line)
        line.replace(r'\**', '')

        if line.strip():
            doc.add_paragraph(line.strip())





    # Add the response text to the document
    # doc.add_paragraph(response_text)

    # temp_dir = tempfile.gettempdir()
    # temp_file = os.path.join(temp_dir, "medical_report.docx")
    # doc.save(temp_file)

    temp_buffer = BytesIO()
    doc.save(temp_buffer)
    temp_buffer.seek(0)

    # Download the document

    # shutil.move("medical_report.docx", f"D:/Atria/Atria/#{patient_id} Health assessment.docx")

    return temp_buffer

    # Send the temporary file as a download
    # return FileResponse(path=temp_file, filename="medical_report.docx",
    #                     media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    # Save the document
    # doc.save("medical_report.docx")

    # Download the document

    # shutil.move("medical_report.docx", f"C:/Users/AbhinavBharti/Documents/atria/#{patient_id} Health assessment.docx")

    # return response.choices[0].message.content





# Function to load data and create the vector database
def load_data_and_create_vector_db(patient_id):


# -------------------- FETCH IMAGE REPORTS--------------------------------

    docs = {}

    cursor = conn.cursor()
    image_results = execute_query(
        cursor,
        f"SELECT * FROM {schema}.ADE_ORDER_RESULT WHERE DOCUMENTCLASS = 'IMAGINGRESULT' AND PATIENTID = '{patient_id}'"
    )

    # Retrieve column names from cursor description
    # Check if there are any results
    if image_results:
        # Retrieve column names from cursor description
        column_names = [desc[0] for desc in cursor.description]

        # Initialize empty lists for each column
        for column_name in column_names:
            docs[column_name] = []

        # Iterate over the result rows
        for result in image_results:
            # Append data to the respective column lists
            for i, column_name in enumerate(column_names):
                docs[column_name].append(result[i])
    else:
        print("No results found.")
    docs_df_report_image = pd.DataFrame(docs)
    docs_df_report_image['source'] = 'Image Report'
    # docs_df_report_image.to_csv('docs_df_report_image.csv', index=False)

    # print(docs_df_report_image)
# -------------------------FETCH LAB REPORTS--------------------------------------------



    observations = execute_query(
        cursor,
        f"SELECT * FROM {schema}.ADE_ORDER_OBSERVATION WHERE DOCUMENTID IN (SELECT DOCUMENTID FROM {schema}.ADE_ORDER_RESULT WHERE DOCUMENTCLASS = 'LABRESULT' AND PATIENTID = '{patient_id}')"
    )


    docs = []

    if observations:
        # Retrieve column names from cursor description
        column_names = [desc[0] for desc in cursor.description]

        # Iterate over the result rows
        for observation in observations:
            # Initialize a dictionary to store observation data
            obs_data = {}

            # Process observation data using descriptive column names
            for column_name, value in zip(column_names, observation):
                obs_data[column_name] = value

            # Append the processed observation data to the list
            docs.append(obs_data)
    else:
        print("No observations found for lab result.")

    # Convert to DataFrame and add source column
    docs_df_report_lab = pd.DataFrame(docs)
    docs_df_report_lab['source'] = 'Lab Report'
    # docs_df_report_lab.to_csv('docs_df_report_lab.csv', index=False)
    # print(docs_df_report_lab)


# -----------------------Fetch Doctors Notes------------------------------------------



    encounter_notes = []
    notes = execute_query(
        cursor,
        f"SELECT ECNOUNTER_NOTE FROM {schema}.ADE_ENCOUNTER_NOTE WHERE PATIENTID = '{patient_id}'"
    )



    if notes:
        # Retrieve column names from cursor description (in this case, we only have one column)
        column_names = ["ECNOUNTER_NOTE"]


        for note in notes:
            # Initialize a dictionary to store note data
            note_data = {}

            # Process note data using descriptive column names
            for column_name, value in zip(column_names, note):
                note_data[column_name] = value


            # Append the processed note data to the list
            encounter_notes.append(note_data)
    else:
        print("No encounter notes found for the patient.")

        # Process note data

    # Convert to DataFrame and add source column
    docs_df_report_notes = pd.DataFrame(encounter_notes)
    docs_df_report_notes['source'] = "Doctor's Note"
    # print(docs_df_report_notes)



    return docs_df_report_image, docs_df_report_lab, docs_df_report_notes


# Mock database of users
users = {
    "satwik": "password1",
}

# Initialize session state
if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False


def check_credentials(username, password):
    if username in users and users[username] == password:
        st.session_state['logged_in'] = True
        st.experimental_rerun()  # Force rerun to update the page
    else:
        st.session_state['logged_in'] = False
        st.error("Username/password is incorrect. Please try again.")


def logout():
    st.session_state['logged_in'] = False
    st.experimental_rerun()


# Login form
if not st.session_state['logged_in']:
    st.title("Login to your account")
    username = st.text_input("Username", key="username")
    password = st.text_input("Password", type="password", key="password")
    if st.button("Login"):
        check_credentials(username, password)
else:
    main_app()
